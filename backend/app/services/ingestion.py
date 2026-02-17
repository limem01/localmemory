import os
import hashlib
import json
import logging
from pathlib import Path
from typing import Optional, List
from datetime import datetime

from sqlalchemy.orm import Session

from app.models.document import Document, DocumentChunk, DocumentStatus, DocumentType
from app.utils.chunking import chunk_text
from app.utils.embeddings import embed_and_store, delete_document_chunks
from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}


def get_document_type(filename: str) -> DocumentType:
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": DocumentType.PDF,
        ".txt": DocumentType.TXT,
        ".docx": DocumentType.DOCX,
        ".md": DocumentType.MD,
    }
    return mapping.get(ext, DocumentType.UNKNOWN)


def compute_file_hash(file_path: str) -> str:
    """Compute SHA-256 hash of a file."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            sha256.update(block)
    return sha256.hexdigest()


def extract_text(file_path: str, doc_type: DocumentType) -> str:
    """Extract plain text from a document."""
    if doc_type == DocumentType.PDF:
        return extract_pdf(file_path)
    elif doc_type == DocumentType.DOCX:
        return extract_docx(file_path)
    elif doc_type in (DocumentType.TXT, DocumentType.MD):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        # Try as plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception:
            raise ValueError(f"Unsupported document type: {doc_type}")


def extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        text_parts = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise


def extract_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise


class IngestionService:
    def __init__(self, db: Session):
        self.db = db

    async def ingest_file(
        self,
        file_path: str,
        original_filename: str,
        title: Optional[str] = None,
        tags: Optional[List[str]] = None,
        is_watched: bool = False,
    ) -> Document:
        """Ingest a file into the system."""
        file_path = str(Path(file_path).resolve())
        filename = Path(original_filename).name
        doc_type = get_document_type(filename)

        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Unsupported file type: {Path(filename).suffix}")

        # Check for duplicates via content hash
        content_hash = compute_file_hash(file_path)
        existing = self.db.query(Document).filter(
            Document.content_hash == content_hash
        ).first()
        if existing:
            logger.info(f"Document already exists: {existing.id} (hash match)")
            return existing

        # Create document record
        doc = Document(
            title=title or Path(filename).stem,
            filename=filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            doc_type=doc_type,
            status=DocumentStatus.PROCESSING,
            is_watched=is_watched,
            content_hash=content_hash,
            tags=json.dumps(tags) if tags else None,
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        try:
            await self._process_document(doc)
        except Exception as e:
            doc.status = DocumentStatus.FAILED
            doc.error_message = str(e)
            self.db.commit()
            logger.error(f"Failed to process document {doc.id}: {e}", exc_info=True)
            raise

        return doc

    async def _process_document(self, doc: Document):
        """Extract text, chunk it, embed and store."""
        logger.info(f"Processing document {doc.id}: {doc.filename}")

        # Extract text
        text = extract_text(doc.file_path, doc.doc_type)
        if not text.strip():
            raise ValueError("Document appears to be empty or unreadable")

        # Chunk text
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("No text chunks generated from document")

        # Store chunks in DB
        db_chunks = []
        for chunk in chunks:
            db_chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=chunk["index"],
                content=chunk["content"],
                word_count=chunk["word_count"],
            )
            self.db.add(db_chunk)
            db_chunks.append(db_chunk)

        self.db.flush()

        # Embed and store in ChromaDB
        chroma_ids = await embed_and_store(
            chunks=chunks,
            document_id=doc.id,
            document_title=doc.title,
        )

        # Update chunk records with chroma IDs
        for db_chunk, chroma_id in zip(db_chunks, chroma_ids):
            db_chunk.chroma_id = chroma_id

        # Update document status
        doc.status = DocumentStatus.READY
        doc.chunk_count = len(chunks)
        self.db.commit()

        logger.info(f"Document {doc.id} processed: {len(chunks)} chunks")

    async def delete_document(self, document_id: int) -> bool:
        """Delete a document and all its data."""
        doc = self.db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return False

        # Delete from ChromaDB
        delete_document_chunks(document_id)

        # Delete file if it exists
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except Exception as e:
                logger.warning(f"Could not delete file {doc.file_path}: {e}")

        # Delete from DB (cascades to chunks)
        self.db.delete(doc)
        self.db.commit()

        return True

    async def ingest_watched_folder(self, folder_path: str):
        """Scan folder and ingest any new files."""
        if not os.path.exists(folder_path):
            return

        for entry in os.scandir(folder_path):
            if not entry.is_file():
                continue

            ext = Path(entry.name).suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue

            # Check if already ingested
            content_hash = compute_file_hash(entry.path)
            existing = self.db.query(Document).filter(
                Document.content_hash == content_hash
            ).first()
            if existing:
                continue

            try:
                logger.info(f"Auto-ingesting watched file: {entry.name}")
                await self.ingest_file(
                    file_path=entry.path,
                    original_filename=entry.name,
                    is_watched=True,
                )
            except Exception as e:
                logger.error(f"Failed to auto-ingest {entry.name}: {e}")
